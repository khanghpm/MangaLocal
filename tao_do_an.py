"""
==============================================================================
  MASTER CODE - TẠO FILE WORD ĐỒ ÁN TỐT NGHIỆP CHUẨN
  Trường Cao Đẳng Bách Khoa Sài Gòn
  Chuẩn: Font TNR 13 | Lề T/P/D: 2cm, Trái: 2.5cm | Giãn dòng 1.5
==============================================================================
  Cài đặt: pip install python-docx
  Chạy:    python tao_do_an.py
  Output:  do_an_tot_nghiep.docx
==============================================================================
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from lxml import etree


# ==============================================================================
# ██████  PHẦN 1: KHAI BÁO NỘI DUNG - CHỈ CẦN CHỈNH SỬA PHẦN NÀY
# ==============================================================================

# --- THÔNG TIN TRANG BÌA ---
BI_TRUONG       = "TRƯỜNG CAO ĐẲNG BÁCH KHOA SÀI GÒN"
BI_KHOA         = "KHOA CÔNG NGHỆ THÔNG TIN & MỸ THUẬT CÔNG NGHIỆP"
BI_LOAI         = "KHOÁ LUẬN TỐT NGHIỆP"
BI_TEN_DE_TAI   = "XÂY DỰNG WEBSITE BÁN HÀNG ĐA NỀN TẢNG OMNIGO"
BI_GVHD         = "Võ Hoàng Tuấn"
BI_SVTH         = [
    "Thân Hoàng Long   -   MSSV: 22000218",
    "Lê Thái Ngọc      -   MSSV: 22001444",
    "Nguyễn Huỳnh Thẳng Thắng  -  MSSV: 22001655",
]
BI_LOP          = "22CPM01H4N"
BI_KHOA_HOC     = "2022"
BI_THANG_NAM    = "Thành phố Hồ Chí Minh, tháng 6 năm 2026"

# --- LỜI CẢM ƠN ---
NOI_DUNG_LOI_CAM_ON = [
    "Trước hết, nhóm chúng em xin gửi lời cảm ơn chân thành và sâu sắc đến Ban Giám hiệu "
    "Trường Cao đẳng Bách khoa Sài Gòn cùng toàn thể quý thầy cô Khoa Công nghệ Thông tin "
    "và Mỹ thuật Công nghiệp đã tận tình giảng dạy, truyền đạt cho chúng em những kiến thức "
    "chuyên môn và kinh nghiệm quý báu trong suốt thời gian học tập tại trường.",

    "Đặc biệt, nhóm chúng em xin bày tỏ lòng biết ơn sâu sắc đến thầy Võ Hoàng Tuấn, người "
    "đã trực tiếp hướng dẫn và đồng hành cùng nhóm trong suốt quá trình thực hiện khóa luận. "
    "Thầy đã tận tình chỉ dẫn, định hướng nội dung, góp ý chi tiết và giúp nhóm chỉnh sửa, "
    "hoàn thiện các phần còn thiếu sót.",

    "Do thời gian thực hiện và kinh nghiệm thực tế còn hạn chế, khóa luận không tránh khỏi "
    "những thiếu sót. Nhóm chúng em kính mong quý thầy cô thông cảm và đóng góp ý kiến để "
    "nhóm có thể tiếp tục hoàn thiện kiến thức và kỹ năng trong tương lai.",

    "Nhóm chúng em xin chân thành cảm ơn.",
]

# Bảng phân công nhiệm vụ (danh sách [STT, Họ tên, MSSV, Nhiệm vụ])
BANG_PHAN_CONG = [
    ["1", "Thân Hoàng Long",              "22000218", "Cung cấp dữ liệu và lập kế hoạch xây dựng sản phẩm"],
    ["2", "Lê Thái Ngọc",                 "22001444", "Xây dựng sản phẩm, cập nhật và tối ưu theo yêu cầu"],
    ["3", "Nguyễn Huỳnh Thẳng Thắng",     "22001655", "Kiểm duyệt, nắm bắt nội dung và truyền đạt"],
]

# --- LỜI MỞ ĐẦU ---
NOI_DUNG_MO_DAU = [
    "Trong bối cảnh thương mại điện tử phát triển mạnh mẽ tại Việt Nam, việc xây dựng một nền "
    "tảng bán hàng trực tuyến tích hợp đa kênh trở nên cấp thiết hơn bao giờ hết. Đề tài "
    "\"Xây dựng website bán hàng đa nền tảng OMNIGO\" ra đời nhằm đáp ứng nhu cầu đó.",

    "Mục tiêu của đề tài là xây dựng hệ thống website bán hàng hoàn chỉnh hỗ trợ quản lý sản "
    "phẩm, đơn hàng, khách hàng và thanh toán trực tuyến. Hệ thống được phát triển bằng "
    "TypeScript, Dart và PostgreSQL (PL/pgSQL).",

    "Bố cục đồ án gồm 3 chương: Chương 1 trình bày tổng quan đề tài; Chương 2 trình bày cơ "
    "sở lý thuyết và công nghệ sử dụng; Chương 3 trình bày phân tích và thiết kế hệ thống.",
]

# --- NỘI DUNG CÁC CHƯƠNG ---
# Cấu trúc: { "tieu_de_chuong": "...", "muc": [ {"title": "...", "level": 1|2, "noi_dung": [...]} ] }

CHUONG_1 = {
    "so_chuong": "CHƯƠNG 1",
    "tieu_de": "TỔNG QUAN ĐỀ TÀI",
    "muc": [
        {
            "title": "1.1 Tính cấp thiết của đề tài",
            "level": 1,
            "noi_dung": [
                "Thương mại điện tử tại Việt Nam tăng trưởng trung bình 25% mỗi năm. Tuy nhiên, "
                "hầu hết các doanh nghiệp vừa và nhỏ vẫn chưa có nền tảng bán hàng trực tuyến "
                "đủ mạnh để cạnh tranh. OMNIGO ra đời nhằm lấp đầy khoảng trống đó.",
            ]
        },
        {
            "title": "1.2 Mục tiêu nghiên cứu",
            "level": 1,
            "noi_dung": [
                "Nghiên cứu nhằm đạt các mục tiêu sau: xây dựng hệ thống quản lý bán hàng "
                "đa nền tảng, tích hợp thanh toán trực tuyến và quản lý kho hàng tự động.",
            ]
        },
        {
            "title": "1.2.1 Mục tiêu tổng quát",
            "level": 2,
            "noi_dung": [
                "Xây dựng website bán hàng đầy đủ chức năng, hoạt động ổn định trên nhiều "
                "thiết bị và trình duyệt khác nhau.",
            ]
        },
        {
            "title": "1.3 Phạm vi nghiên cứu",
            "level": 1,
            "noi_dung": [
                "Đề tài tập trung vào phân hệ dành cho khách hàng (đặt hàng, thanh toán) "
                "và phân hệ dành cho Admin (quản lý đơn, sản phẩm, khách hàng).",
            ]
        },
    ]
}

CHUONG_2 = {
    "so_chuong": "CHƯƠNG 2",
    "tieu_de": "CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ",
    "muc": [
        {
            "title": "2.1 Tổng quan website bán hàng",
            "level": 1,
            "noi_dung": [
                "Website bán hàng (e-commerce) là hệ thống cho phép giao dịch mua bán hàng hóa "
                "và dịch vụ qua mạng Internet. Hệ thống bao gồm giao diện người dùng, hệ thống "
                "thanh toán và cơ sở dữ liệu quản lý sản phẩm.",
            ]
        },
        {
            "title": "2.2 Tổng quan về Dart",
            "level": 1,
            "noi_dung": [
                "Dart là ngôn ngữ lập trình được Google phát triển, hỗ trợ lập trình hướng đối "
                "tượng và bất đồng bộ. Trong dự án này, Dart được sử dụng ở phía server.",
            ]
        },
        {
            "title": "2.2.1 Giới thiệu chung về Dart",
            "level": 2,
            "noi_dung": [
                "Dart được thiết kế để tối ưu cho việc xây dựng ứng dụng client và server, "
                "với cú pháp gần gũi với Java và C#. Dart hỗ trợ null safety từ phiên bản 2.12.",
            ]
        },
        {
            "title": "2.3 Tổng quan về PostgreSQL",
            "level": 1,
            "noi_dung": [
                "PostgreSQL là hệ quản trị cơ sở dữ liệu quan hệ mã nguồn mở mạnh mẽ, hỗ trợ "
                "ACID, stored procedures (PL/pgSQL) và JSON. Đây là lựa chọn phù hợp cho hệ "
                "thống thương mại điện tử đòi hỏi tính nhất quán dữ liệu cao.",
            ]
        },
    ]
}

CHUONG_3 = {
    "so_chuong": "CHƯƠNG 3",
    "tieu_de": "PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG",
    "muc": [
        {
            "title": "3.1 Phân tích yêu cầu hệ thống",
            "level": 1,
            "noi_dung": [
                "Hệ thống cần đáp ứng hai nhóm yêu cầu chính: yêu cầu chức năng (functional "
                "requirements) và yêu cầu phi chức năng (non-functional requirements).",
            ]
        },
        {
            "title": "3.1.1 Yêu cầu chức năng",
            "level": 2,
            "noi_dung": [
                "Phân hệ khách hàng: đăng ký/đăng nhập, xem sản phẩm, thêm vào giỏ hàng, "
                "đặt hàng và thanh toán. Phân hệ Admin: quản lý sản phẩm, đơn hàng, người dùng.",
            ]
        },
        {
            "title": "3.2 Thiết kế cơ sở dữ liệu",
            "level": 1,
            "noi_dung": [
                "Cơ sở dữ liệu được thiết kế theo mô hình quan hệ với các bảng chính: users, "
                "products, categories, orders, order_items, payments. Mô hình ERD thể hiện "
                "đầy đủ các mối quan hệ giữa các thực thể.",
            ]
        },
    ]
}

# --- KẾT LUẬN ---
NOI_DUNG_KET_LUAN = [
    "Sau quá trình nghiên cứu và thực hiện, nhóm đã hoàn thành xây dựng website bán hàng đa "
    "nền tảng OMNIGO với đầy đủ chức năng theo yêu cầu đặt ra ban đầu.",

    "Kết quả đạt được: hệ thống hoạt động ổn định, giao diện thân thiện với người dùng, "
    "tích hợp thanh toán trực tuyến và quản lý đơn hàng hiệu quả.",

    "Hạn chế và hướng phát triển: nhóm dự kiến mở rộng tính năng gợi ý sản phẩm bằng AI, "
    "tích hợp thêm cổng thanh toán và phát triển ứng dụng di động trong tương lai.",
]

# --- TÀI LIỆU THAM KHẢO ---
TAI_LIEU_THAM_KHAO = [
    "Lê Văn A. 2023. Lập trình web với TypeScript. Hà Nội: Nhà xuất bản Khoa học Kỹ thuật.",
    "PostgreSQL Documentation. 2024. PL/pgSQL – SQL Procedural Language. postgresql.org.",
    "Google LLC. 2023. Dart Programming Language Documentation. dart.dev.",
    "VECOM. 2024. Báo cáo thương mại điện tử Việt Nam 2024. Hà Nội: VECOM.",
]


# ==============================================================================
# ██████  PHẦN 2: ENGINE TẠO VĂN BẢN - KHÔNG CẦN CHỈNH SỬA
# ==============================================================================

def _set_font_xml(run, font_name="Times New Roman", size_pt=13):
    """Ép font Times New Roman qua XML để đảm bảo cả tiếng Việt có dấu."""
    rPr = run._r.get_or_add_rPr()
    # Xóa rFonts cũ nếu có
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),       font_name)
    rFonts.set(qn('w:hAnsi'),       font_name)
    rFonts.set(qn('w:cs'),          font_name)  # Critical cho tiếng Việt
    rFonts.set(qn('w:eastAsia'),    font_name)
    rPr.insert(0, rFonts)
    run.font.size = Pt(size_pt)


def _set_paragraph_format(para, first_line_cm=1.0, space_after_pt=0, justify=True):
    """Thiết lập định dạng chuẩn cho đoạn văn."""
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = 1.5
    pf.first_line_indent = Cm(first_line_cm)
    pf.space_after       = Pt(space_after_pt)
    pf.space_before      = Pt(0)
    if justify:
        pf.alignment     = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_body_text(doc, text, first_indent=True, space_after=6):
    """Thêm đoạn văn bản thường với định dạng chuẩn."""
    para = doc.add_paragraph()
    run  = para.add_run(text)
    _set_font_xml(run, size_pt=13)
    _set_paragraph_format(para,
                          first_line_cm=1.0 if first_indent else 0,
                          space_after_pt=space_after)
    return para


def add_heading_chuong(doc, so_chuong, tieu_de):
    """Thêm tiêu đề CHƯƠNG: in hoa, đậm, cỡ 14, canh giữa."""
    # Dòng số chương
    p1 = doc.add_paragraph()
    r1 = p1.add_run(so_chuong)
    _set_font_xml(r1, size_pt=14)
    r1.bold = True
    p1.paragraph_format.alignment     = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before  = Pt(0)
    p1.paragraph_format.space_after   = Pt(0)
    p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p1.paragraph_format.line_spacing  = 1.5
    p1.paragraph_format.first_line_indent = Cm(0)

    # Dòng tên chương
    p2 = doc.add_paragraph()
    r2 = p2.add_run(tieu_de)
    _set_font_xml(r2, size_pt=14)
    r2.bold = True
    p2.paragraph_format.alignment     = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before  = Pt(0)
    p2.paragraph_format.space_after   = Pt(12)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p2.paragraph_format.line_spacing  = 1.5
    p2.paragraph_format.first_line_indent = Cm(0)
    return p2


def add_heading_muc(doc, text, level=1):
    """
    Thêm tiêu đề mục con.
    level=1 → Mục lớn (1.1): đậm, size 13, lề trái
    level=2 → Tiểu mục (1.1.1): đậm nghiêng, size 13, lề trái
    """
    para = doc.add_paragraph()
    run  = para.add_run(text)
    _set_font_xml(run, size_pt=13)
    run.bold   = True
    run.italic = (level == 2)
    pf = para.paragraph_format
    pf.alignment          = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent  = Cm(0)
    pf.space_before       = Pt(6)
    pf.space_after        = Pt(3)
    pf.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing       = 1.5
    return para


def add_page_break(doc):
    """Thêm ngắt trang."""
    doc.add_page_break()


def add_section_title(doc, text):
    """Tiêu đề phần lớn (LỜI CẢM ƠN, MỤC LỤC, KẾT LUẬN...): in hoa, đậm, giữa, size 14."""
    para = doc.add_paragraph()
    run  = para.add_run(text.upper())
    _set_font_xml(run, size_pt=14)
    run.bold = True
    pf = para.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = 1.5
    pf.first_line_indent = Cm(0)
    return para


def add_table(doc, headers, rows, col_widths_cm=None):
    """
    Thêm bảng có viền chuẩn.
    headers: list tên cột
    rows: list of list dữ liệu
    col_widths_cm: list chiều rộng cột (cm), nếu None thì chia đều
    """
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'

    # Tổng chiều rộng khả dụng: 21 - 2.5 - 2 = 16.5 cm
    total_cm    = 16.5
    if col_widths_cm is None:
        col_widths_cm = [total_cm / n_cols] * n_cols

    # Header row
    hdr_row = table.rows[0]
    for i, htext in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = Cm(col_widths_cm[i])
        para = cell.paragraphs[0]
        run  = para.add_run(htext)
        _set_font_xml(run, size_pt=13)
        run.bold = True
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_obj = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row_obj.cells[c_idx]
            cell.width = Cm(col_widths_cm[c_idx])
            para = cell.paragraphs[0]
            run  = para.add_run(str(cell_text))
            _set_font_xml(run, size_pt=13)
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
    return table


def add_blank_line(doc):
    """Thêm dòng trắng (cách đoạn)."""
    para = doc.add_paragraph()
    run  = para.add_run("")
    _set_font_xml(run, size_pt=13)
    para.paragraph_format.space_before  = Pt(0)
    para.paragraph_format.space_after   = Pt(0)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing  = 1.5


# ==============================================================================
# ██████  PHẦN 3: CÁC TRANG ĐẶC BIỆT
# ==============================================================================

def tao_trang_bia(doc, la_phu_bia=False):
    """Tạo trang bìa chính hoặc bìa phụ dùng bảng ẩn để căn chỉnh."""

    def _add_center_bold(doc, text, size, space_after=6, italic=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        _set_font_xml(r, size_pt=size)
        r.bold   = True
        r.italic = italic
        p.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before       = Pt(0)
        p.paragraph_format.space_after        = Pt(space_after)
        p.paragraph_format.first_line_indent  = Cm(0)
        p.paragraph_format.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing       = 1.5
        return p

    # Header trường
    _add_center_bold(doc, "BỘ GIÁO DỤC VÀ ĐÀO TẠO", 13, space_after=0)
    _add_center_bold(doc, BI_TRUONG, 13, space_after=0)
    _add_center_bold(doc, BI_KHOA, 13, space_after=24)

    # ========== DÒNG KẺ ========== (dùng paragraph với border dưới)
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after  = Pt(18)
    pPr   = p_line._p.get_or_add_pPr()
    pBdr  = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    _add_center_bold(doc, BI_LOAI, 16, space_after=18)

    # Tên đề tài - nổi bật nhất
    p_ten = doc.add_paragraph()
    r_ten = p_ten.add_run(BI_TEN_DE_TAI)
    _set_font_xml(r_ten, size_pt=18)
    r_ten.bold = True
    p_ten.paragraph_format.alignment        = WD_ALIGN_PARAGRAPH.CENTER
    p_ten.paragraph_format.space_before     = Pt(0)
    p_ten.paragraph_format.space_after      = Pt(36)
    p_ten.paragraph_format.first_line_indent = Cm(0)
    p_ten.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p_ten.paragraph_format.line_spacing     = 1.5

    # Bảng ẩn: GVHD / SVTH
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Ẩn toàn bộ viền
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)

    def _add_bi_row(table, label, value):
        row   = table.add_row()
        c_lbl = row.cells[0]
        c_val = row.cells[1]
        c_lbl.width = Cm(5.5)
        c_val.width = Cm(11.0)

        p_lbl = c_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        _set_font_xml(r_lbl, size_pt=13)
        r_lbl.bold = True
        p_lbl.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_lbl.paragraph_format.first_line_indent = Cm(0)
        p_lbl.paragraph_format.space_before = Pt(2)
        p_lbl.paragraph_format.space_after  = Pt(2)
        p_lbl.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p_lbl.paragraph_format.line_spacing = 1.5

        p_val = c_val.paragraphs[0]
        r_val = p_val.add_run(value)
        _set_font_xml(r_val, size_pt=13)
        r_val.bold = True
        p_val.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_val.paragraph_format.first_line_indent = Cm(0)
        p_val.paragraph_format.space_before = Pt(2)
        p_val.paragraph_format.space_after  = Pt(2)
        p_val.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p_val.paragraph_format.line_spacing = 1.5

    _add_bi_row(table, "Giảng viên hướng dẫn:", BI_GVHD)
    _add_bi_row(table, "Sinh viên thực hiện:",  BI_SVTH[0])
    for svth in BI_SVTH[1:]:
        _add_bi_row(table, "", svth)
    _add_bi_row(table, "Lớp:",  BI_LOP)
    _add_bi_row(table, "Khoá:", BI_KHOA_HOC)

    # Thành phố & năm
    p_tp = doc.add_paragraph()
    r_tp = p_tp.add_run(BI_THANG_NAM)
    _set_font_xml(r_tp, size_pt=13)
    r_tp.bold = True
    p_tp.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    p_tp.paragraph_format.space_before      = Pt(36)
    p_tp.paragraph_format.space_after       = Pt(0)
    p_tp.paragraph_format.first_line_indent = Cm(0)
    p_tp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p_tp.paragraph_format.line_spacing      = 1.5


def tao_loi_cam_on(doc):
    add_section_title(doc, "LỜI CẢM ƠN")
    for doan in NOI_DUNG_LOI_CAM_ON:
        add_body_text(doc, doan)
    add_blank_line(doc)
    # Bảng phân công
    add_section_title(doc, "BẢNG PHÂN CÔNG NHIỆM VỤ")
    add_table(
        doc,
        headers=["STT", "Họ và tên", "MSSV", "Nhiệm vụ"],
        rows=BANG_PHAN_CONG,
        col_widths_cm=[1.0, 4.5, 3.0, 8.0]
    )


def tao_muc_luc(doc):
    add_section_title(doc, "MỤC LỤC")
    muc_luc_items = [
        ("LỜI CẢM ƠN",                          "I"),
        ("MỤC LỤC",                              "II"),
        ("LỜI MỞ ĐẦU",                           "III"),
        ("CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI",           "1"),
        ("    1.1  Tính cấp thiết của đề tài",    "1"),
        ("    1.2  Mục tiêu nghiên cứu",          "2"),
        ("        1.2.1  Mục tiêu tổng quát",     "2"),
        ("    1.3  Phạm vi nghiên cứu",           "2"),
        ("CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ", "7"),
        ("    2.1  Tổng quan website bán hàng",   "7"),
        ("    2.2  Tổng quan về Dart",            "12"),
        ("        2.2.1  Giới thiệu chung về Dart","12"),
        ("    2.3  Tổng quan về PostgreSQL",      "16"),
        ("CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", "28"),
        ("    3.1  Phân tích yêu cầu hệ thống",  "28"),
        ("        3.1.1  Yêu cầu chức năng",     "28"),
        ("    3.2  Thiết kế cơ sở dữ liệu",      "53"),
        ("KẾT LUẬN",                              "XX"),
        ("TÀI LIỆU THAM KHẢO",                   "XX"),
    ]
    # Bảng mục lục (ẩn viền, 2 cột: nội dung | số trang)
    tbl = doc.add_table(rows=len(muc_luc_items), cols=2)
    tblPr_el = tbl._tbl.find(qn('w:tblPr'))
    if tblPr_el is None:
        tblPr_el = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr_el)
    tblBdr = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:val'), 'none')
        tblBdr.append(e)
    tblPr_el.append(tblBdr)

    for idx, (text, page) in enumerate(muc_luc_items):
        row = tbl.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Cm(14.0)
        c1.width = Cm(2.5)
        is_chapter = text.startswith("CHƯƠNG") or text in ("LỜI CẢM ƠN","MỤC LỤC","LỜI MỞ ĐẦU","KẾT LUẬN","TÀI LIỆU THAM KHẢO")
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(text)
        _set_font_xml(r0, size_pt=13)
        r0.bold = is_chapter
        p0.paragraph_format.first_line_indent = Cm(0)
        p0.paragraph_format.space_before = Pt(1)
        p0.paragraph_format.space_after  = Pt(1)
        p0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p0.paragraph_format.line_spacing = 1.5
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(page)
        _set_font_xml(r1, size_pt=13)
        r1.bold = is_chapter
        p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1.paragraph_format.first_line_indent = Cm(0)
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after  = Pt(1)
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p1.paragraph_format.line_spacing = 1.5


def tao_mo_dau(doc):
    add_section_title(doc, "LỜI MỞ ĐẦU")
    for doan in NOI_DUNG_MO_DAU:
        add_body_text(doc, doan)


def tao_chuong(doc, chuong_data):
    add_heading_chuong(doc, chuong_data["so_chuong"], chuong_data["tieu_de"])
    for muc in chuong_data["muc"]:
        add_heading_muc(doc, muc["title"], level=muc["level"])
        for doan in muc["noi_dung"]:
            add_body_text(doc, doan)
        add_blank_line(doc)
    # Tóm tắt chương
    p_tt = doc.add_paragraph()
    r_tt = p_tt.add_run(f"→ Tóm tắt {chuong_data['so_chuong'].lower()}: "
                        f"Chương này đã trình bày {chuong_data['tieu_de'].lower()}. "
                        "Các nội dung trên là nền tảng để tiếp tục chương tiếp theo.")
    _set_font_xml(r_tt, size_pt=13)
    r_tt.italic = True
    p_tt.paragraph_format.alignment        = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_tt.paragraph_format.first_line_indent = Cm(1.0)
    p_tt.paragraph_format.space_before     = Pt(6)
    p_tt.paragraph_format.space_after      = Pt(0)
    p_tt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p_tt.paragraph_format.line_spacing     = 1.5


def tao_ket_luan(doc):
    add_section_title(doc, "KẾT LUẬN")
    for doan in NOI_DUNG_KET_LUAN:
        add_body_text(doc, doan)


def tao_tai_lieu_tham_khao(doc):
    add_section_title(doc, "TÀI LIỆU THAM KHẢO")
    for item in sorted(TAI_LIEU_THAM_KHAO):
        p = doc.add_paragraph()
        r = p.add_run(f"- {item}")
        _set_font_xml(r, size_pt=13)
        p.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent  = Cm(0)
        p.paragraph_format.left_indent        = Cm(0.5)
        p.paragraph_format.space_before       = Pt(3)
        p.paragraph_format.space_after        = Pt(3)
        p.paragraph_format.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing       = 1.5


# ==============================================================================
# ██████  PHẦN 4: THIẾT LẬP TRANG & XUẤT FILE
# ==============================================================================

def setup_page_margins(doc):
    """Thiết lập lề trang theo chuẩn: Trái 2.5cm, T/P/D 2.0cm."""
    for section in doc.sections:
        section.page_width   = Cm(21)    # A4
        section.page_height  = Cm(29.7)  # A4
        section.left_margin  = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.top_margin   = Cm(2.0)
        section.bottom_margin = Cm(2.0)


def set_default_font_xml(doc):
    """Ép font mặc định TNR 13 vào docDefaults qua XML."""
    styles_elem = doc.styles.element
    docDefaults = styles_elem.find(qn('w:docDefaults'))
    if docDefaults is None:
        return
    rPrDefault = docDefaults.find('.//' + qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = OxmlElement('w:rPrDefault')
        docDefaults.append(rPrDefault)
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        rPrDefault.append(rPr)
    # Xóa rFonts cũ
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    'Times New Roman')
    rFonts.set(qn('w:hAnsi'),    'Times New Roman')
    rFonts.set(qn('w:cs'),       'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rPr.insert(0, rFonts)
    # Cỡ chữ mặc định
    for old_sz in rPr.findall(qn('w:sz')):
        rPr.remove(old_sz)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '26')   # 13pt = 26 half-points
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '26')
    rPr.append(szCs)


def main():
    doc = Document()
    setup_page_margins(doc)
    set_default_font_xml(doc)

    # ── TRANG BÌA CHÍNH ────────────────────────────────────────────────
    tao_trang_bia(doc)
    add_page_break(doc)

    # ── TRANG PHỤ BÌA ──────────────────────────────────────────────────
    tao_trang_bia(doc, la_phu_bia=True)
    add_page_break(doc)

    # ── LỜI CẢM ƠN + PHÂN CÔNG ─────────────────────────────────────────
    tao_loi_cam_on(doc)
    add_page_break(doc)

    # ── MỤC LỤC ────────────────────────────────────────────────────────
    tao_muc_luc(doc)
    add_page_break(doc)

    # ── LỜI MỞ ĐẦU ─────────────────────────────────────────────────────
    tao_mo_dau(doc)
    add_page_break(doc)

    # ── CHƯƠNG 1 ────────────────────────────────────────────────────────
    tao_chuong(doc, CHUONG_1)
    add_page_break(doc)

    # ── CHƯƠNG 2 ────────────────────────────────────────────────────────
    tao_chuong(doc, CHUONG_2)
    add_page_break(doc)

    # ── CHƯƠNG 3 ────────────────────────────────────────────────────────
    tao_chuong(doc, CHUONG_3)
    add_page_break(doc)

    # ── KẾT LUẬN ────────────────────────────────────────────────────────
    tao_ket_luan(doc)
    add_page_break(doc)

    # ── TÀI LIỆU THAM KHẢO ─────────────────────────────────────────────
    tao_tai_lieu_tham_khao(doc)

    # ── XUẤT FILE ───────────────────────────────────────────────────────
    output_path = "do_an_tot_nghiep.docx"
    doc.save(output_path)
    print(f"✅ Xuất thành công: {output_path}")
    print("   → Font: Times New Roman 13pt (ép qua XML)")
    print("   → Lề: Trái 2.5cm | Trên/Phải/Dưới 2.0cm")
    print("   → Giãn dòng: 1.5 | Thụt đầu dòng: 1.0cm")
    print("   → Ngắt trang: sau mỗi phần lớn")


if __name__ == "__main__":
    main()
